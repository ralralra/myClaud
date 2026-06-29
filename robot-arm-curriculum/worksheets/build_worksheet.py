# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 기본 폰트(맑은 고딕)
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def shade(cell, color="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),color)
    tcPr.append(sh)

def h(text, size=15, color="1F4E79", before=10, after=4):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    return p

def sub(text, color="2E74B5"):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True
    r.font.size = Pt(11.5); r.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    return p

def body(text):
    p = doc.add_paragraph(text); p.paragraph_format.space_after = Pt(2); return p

def mission_box(text):
    t = doc.add_table(rows=1, cols=1); t.style='Table Grid'
    c = t.cell(0,0); shade(c, "FCE4D6")
    p = c.paragraphs[0]; r = p.add_run("🏆 도전 미션  "); r.bold=True; r.font.color.rgb=RGBColor.from_string("C45911")
    r2 = p.add_run(text); r2.font.size=Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def table(headers, rows, widths=None, header_color="D9E2F3"):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'
    hdr = t.rows[0].cells
    for i,htext in enumerate(headers):
        hdr[i].text=''
        p=hdr[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(htext); r.bold=True; r.font.size=Pt(10)
        shade(hdr[i], header_color)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=''
            p=cells[i].paragraphs[0]; r=p.add_run(val); r.font.size=Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# 빈칸 n개
def blanks(n): return ['']*n

# ============ 표지 ============
title = doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=title.add_run("아두이노 6자유도 로봇팔\n학습활동지 (12차시)")
r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor.from_string("1F4E79")
sgt = doc.add_paragraph(); sgt.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sgt.add_run("6자유도 로봇팔로 배우는 스마트 자동화 미션"); r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string("2E74B5")
doc.add_paragraph()
table(["항목","내용","항목","내용"],
      [["수업명","로봇팔 스마트 자동화 미션","역할","조작 / 코드 / 기록 / 발표"],
       ["이름","","팀명",""],
       ["학번/반","","로봇팔 번호",""]])
doc.add_paragraph()
note=doc.add_paragraph(); rr=note.add_run("※ 차시별로 한 장씩 배부하여 사용하고, 11~12차시에는 최종 보고서·자기평가표를 함께 제출한다.")
rr.italic=True; rr.font.size=Pt(9); rr.font.color.rgb=RGBColor.from_string("808080")
doc.add_page_break()

# ============ 차시 데이터 ============
sessions = []

sessions.append(dict(
  no=1, title="로봇팔 깨우기", goal="로봇팔을 조립·배선 점검하고, 조이스틱으로 자유롭게 조작해 본다.",
  mission="전원을 켜고 조이스틱만으로 30초 안에 그리퍼로 '목표 지점'을 터치하라! (선생님이 정한 컵·스티커 등)",
  code="01_joystick_control.ino",
  tables=[("1. 로봇팔 관절 매핑표 (조작하며 채우기)",
           ["서보 번호","연결 핀","움직이는 부위","움직임 설명","확인"],
           [[f"{i}번 서보","","","","□"] for i in range(1,7)]),
          ("2. 워밍업 게임 기록",
           ["시도","성공까지 걸린 시간","가장 어려웠던 점"],
           [["1차",""," "],["2차",""," "],["3차",""," "]])],
  qs=["자동차 로봇과 로봇팔의 가장 큰 차이는 무엇인가?",
      "로봇팔에서 서보모터가 중요한 이유는 무엇인가?",
      "GND(접지) 연결이 중요한 이유는 무엇인가?"]))

sessions.append(dict(
  no=2, title="코드 해부 — 코드 탐정", goal="샘플 코드의 구조(setup/loop)·핀·각도·조건문을 찾고, 값을 바꿔 동작 변화를 관찰한다.",
  mission="코드의 숫자를 일부러 바꿔 무슨 일이 일어나는지 알아내라! (예: delay 줄이기, MAX 줄이기, PIN 바꾸기)",
  code="01_joystick_control.ino",
  tables=[("1. 코드 해석표 (코드에서 값 찾기)",
           ["코드 요소","코드에서 찾은 값","의미(내 말로)"],
           [["서보모터 개수","",""],["서보 핀 번호","",""],["조이스틱 입력 핀","",""],
            ["초기 각도값","",""],["최소/최대 각도","",""],["조이스틱 기준값(612/412)","",""]]),
          ("2. 코드 탐정 실험 — 바꾸고 관찰하기",
           ["바꾼 부분","원래 값 → 바꾼 값","로봇팔에 생긴 변화"],
           [["delay()","",""],["MAX 각도","",""],["PIN(핀 번호)","",""]])],
  qs=["핀 번호를 잘못 쓰면 어떤 문제가 생길까?",
      "612와 412 사이를 '멈춤'으로 둔 이유는 무엇일까?",
      "최소·최대 각도 제한이 필요한 이유는 무엇인가?"]))

sessions.append(dict(
  no=3, title="조이스틱 마스터", goal="조이스틱 정밀 조작 능력을 기르고, 안정적으로 움직이는 방법을 찾는다.",
  mission="[택1] ① 블록 타워 가장 높이 쌓기  ② 컵 옮기기 타임어택 — 팀 최고기록에 도전!",
  code="01_joystick_control.ino",
  tables=[("1. 정밀 조작 도전 기록",
           ["시도","기록(높이/시간)","실패 원인","개선 아이디어"],
           [["1차","","",""],["2차","","",""],["3차","","",""]]),
          ("2. 움직임 관찰",
           ["관찰 항목","기록"],
           [["가장 조작하기 쉬운 관절",""],["가장 어려운 관절",""],
            ["팔이 흔들렸던 상황",""],["안정적으로 움직이는 나만의 요령",""]])],
  qs=["로봇팔을 한 번에 크게 움직이면 어떤 문제가 생길까?",
      "정밀 제어가 필요한 이유는 무엇인가?"]))

sessions.append(dict(
  no=4, title="홈 포지션 & 안전", goal="안정적인 기본 자세(홈 포지션)를 찾아 초기 각도값을 수정한다.",
  mission="1mm 챌린지 — 펜 끝을 매번 '같은 점'에 닿게 하라. 5번 중 몇 번이나 같은 자리에 닿는가?",
  code="01 (INITANGLE 값 수정)",
  tables=[("1. 홈 포지션 각도표",
           ["서보 번호","기존 초기값","수정한 초기값","수정 이유"],
           [[f"{i}번 서보","","",""] for i in range(1,7)]),
          ("2. 안정성 점검",
           ["점검 항목","예/아니오","설명"],
           [["전원을 켰을 때 갑자기 크게 움직이지 않는다","□ 예 / □ 아니오",""],
            ["그리퍼가 바닥에 닿지 않는다","□ 예 / □ 아니오",""],
            ["팔이 너무 앞으로 기울지 않는다","□ 예 / □ 아니오",""],
            ["다시 실행해도 같은 자세로 온다","□ 예 / □ 아니오",""]])],
  qs=["홈 포지션이 필요한 이유는 무엇인가?",
      "내가 수정한 각도값 중 가장 중요한 값은 무엇인가?",
      "안전한 로봇팔 제어를 위해 주의할 점은 무엇인가?"]))

sessions.append(dict(
  no=5, title="그리퍼 장인", goal="그리퍼 각도를 조절해 다양한 물체를 안정적으로(부드럽게) 집는다.",
  mission="깨지기 쉬운 화물 운반 — 마시멜로/종이컵을 터뜨리지 않고 옮겨라! 가장 적은 힘으로 집은 팀 승!",
  code="01_joystick_control.ino",
  tables=[("1. 물체별 집기 실험 기록표",
           ["물체","열린 각도","닫힌 각도","성공 여부","실패 원인"],
           [["스펀지 큐브","","","□ 성공 / □ 실패",""],
            ["종이컵","","","□ 성공 / □ 실패",""],
            ["마시멜로","","","□ 성공 / □ 실패",""],
            ["둥근 물체","","","□ 성공 / □ 실패",""]]),
          ("2. 가장 쉬운/어려운 물체",
           ["항목","내용"],
           [["가장 잘 집힌 물체",""],["이유",""],["가장 어려운 물체",""],["이유",""]])],
  qs=["그리퍼 각도를 너무 작게/크게 하면 어떤 문제가 생기는가?",
      "로봇팔이 물체를 집을 때 고려해야 할 점은 무엇인가?"]))

sessions.append(dict(
  no=6, title="로봇과 대화하기 (Serial)", goal="코드에 Serial을 추가해 현재 각도를 화면으로 읽고, 멋진 자세를 각도값으로 캡처한다.",
  mission="자세 사진사 — 멋진 포즈를 잡고 시리얼 모니터에 'p'를 입력해 각도값을 캡처하라. 친구와 값을 교환해 똑같이 재현해 보기!",
  code="02_serial_record.ino",
  tables=[("1. 캡처한 자세 기록 ( p 입력으로 출력된 값 )",
           ["자세 이름","1번","2번","3번","4번","5번","6번(그리퍼)"],
           [["예) 인사 자세","","","","","",""],["", "","","","","",""],["", "","","","","",""]]),
          ("2. 친구 자세 재현 결과",
           ["받은 자세","재현 성공?","달랐던 점"],
           [["",  "□ 성공 / □ 실패",""]])],
  qs=["로봇의 움직임을 '숫자(각도)'로 기록한다는 것은 어떤 의미인가?",
      "Serial(시리얼)이 없으면 왜 각도값을 기록하기 어려운가?"]))

sessions.append(dict(
  no=7, title="Pick & Place 작전", goal="물체를 A→B로 옮기는 작업 순서를 설계하고 수동으로 수행한다.",
  mission="택배 분류 타임어택 — A구역 물체를 B구역으로! 작업 순서를 먼저 설계하고 가장 빠르고 정확하게.",
  code="02_serial_record.ino",
  tables=[("1. 작업 순서 설계표",
           ["순서","동작","주의할 점"],
           [[str(i),"",""] for i in range(1,7)]),
          ("2. 수행 기록",
           ["시도","걸린 시간","성공 여부","개선점"],
           [["1차","","□ 성공 / □ 실패",""],["2차","","□ 성공 / □ 실패",""]])],
  qs=["작업을 순서대로 나누는 것이 왜 중요한가?",
      "사람이 하던 분류를 로봇이 하면 무엇이 좋아지는가?"]))

sessions.append(dict(
  no=8, title="자세 레시피 북", goal="성공한 자세들을 각도값으로 기록·정리해, 자동화에 쓸 '포즈 카드'를 만든다.",
  mission="HOME·PICK·PLACE 포즈 카드를 완성하라. 이 값이 다음 시간 '자동화'의 재료가 된다!",
  code="02_serial_record.ino",
  tables=[("1. 포즈 카드 (자동화에 넣을 핵심 자세)",
           ["포즈 이름","1번","2번","3번","4번","5번","6번"],
           [["HOME (기본)","","","","","",""],
            ["PICK_READY (집기 직전)","","","","","",""],
            ["PICK (집는 위치)","","","","","",""],
            ["PLACE_READY (놓기 직전)","","","","","",""],
            ["PLACE (놓는 위치)","","","","","",""]]),
          ("2. 점검",
           ["점검 항목","예/아니오"],
           [["같은 값을 다시 입력하면 비슷한 자세가 재현된다","□ 예 / □ 아니오"],
            ["불필요한(중복) 자세를 제외했다","□ 예 / □ 아니오"]])],
  qs=["각도값을 다시 입력했을 때 같은 동작이 나와야 하는 이유는?",
      "기록한 자세 중 가장 중요한 자세는 무엇이고 왜인가?"]))

sessions.append(dict(
  no=9, title="함수 만들기", goal="반복되는 로봇팔 동작을 함수로 만들고, 호출해서 실행한다.",
  mission="내 함수 만들기 — gripperOpen()/gripperClose()부터 만들고, goHome()과 나만의 동작 함수(예: wave())까지 완성하라!",
  code="03_functions_button.ino",
  tables=[("1. 함수 설계표",
           ["함수 이름","하는 일","사용한 명령/값"],
           [["gripperOpen()","그리퍼 열기",""],
            ["gripperClose()","그리퍼 닫기",""],
            ["goHome()","홈 자세로 이동",""],
            ["내가 만든 함수 (        )","",""]]),
          ("2. 함수 전/후 비교",
           ["구분","코드가 어떻게 달라졌나"],
           [["함수 없이 직접 쓸 때(복붙)",""],["함수로 만든 뒤",""]])],
  qs=["함수를 쓰면 무엇이 편리해지는가?",
      "매개변수(괄호 안 값)를 받는 함수의 장점은 무엇인가?"]))

sessions.append(dict(
  no=10, title="자동 Pick & Place", goal="기록한 포즈와 함수를 활용해, 명령(또는 버튼) 한 번으로 자동 동작을 실행한다.",
  mission="원클릭 공장 — 시리얼에 '1'을 입력(또는 버튼)하면 집기→옮기기→놓기가 자동으로! 10번 중 몇 번 성공하나?",
  code="03_functions_button.ino",
  tables=[("1. 자동화 동작 흐름도 (함수 호출 순서)",
           ["순서","호출한 함수","목표 포즈"],
           [[str(i),"",""] for i in range(1,8)]),
          ("2. 자동화 성공률 테스트",
           ["시도 묶음","성공/시도","주된 실패 원인","조정한 값"],
           [["1~10회","   / 10","",""],["11~20회","   / 10","",""]])],
  qs=["수동 조작과 자동 실행의 차이는 무엇인가?",
      "자동 동작이 실패할 때, 코드에서 어떤 값을 조정해야 하나?"]))

sessions.append(dict(
  no=11, title="최종 미션 설계 & 튜닝", goal="팀별 미니 챌린지를 직접 설계하고, 반복 테스트로 성공률을 높인다.",
  mission="우리 팀 미션을 설계하라! (예: 색깔 분류, 컵 피라미드, 도미노, 캘리그래피, 미니 농구…) 수동+자동을 섞어도 좋다.",
  code="03_functions_button.ino",
  tables=[("1. 최종 미션 설계표",
           ["항목","내용"],
           [["미션 이름",""],["목표(성공 기준)",""],["사용 물체/소품",""],
            ["수동/자동 구성",""],["예상 어려움",""]]),
          ("2. 튜닝 기록",
           ["테스트","성공률","바꾼 값","결과"],
           [["1차","","",""],["2차","","",""],["3차","","",""]])],
  qs=["성공률을 높이기 위해 가장 효과적이었던 조정은 무엇인가?",
      "우리 미션을 실제 산업현장과 연결하면 어떤 사례와 닮았는가?"]))

sessions.append(dict(
  no=12, title="데모데이 & 진로", goal="최종 미션을 시연하고, 활동 경험을 로봇·자동화 진로와 연결해 발표한다.",
  mission="데모데이! 다른 팀 앞에서 미션을 시연하고, 우리 로봇팔의 '직업'을 상상해 발표하라.",
  code="03_functions_button.ino",
  tables=[("1. 최종 미션 결과 기록",
           ["항목","내용"],
           [["미션 이름",""],["최종 성공률",""],["가장 자랑스러운 점",""],["다음에 개선할 점",""]]),
          ("2. 진로 연결 발표 준비",
           ["질문","우리 팀 답변"],
           [["이 로봇팔과 닮은 직업/산업은?",""],
            ["오늘 배운 것 중 진로에 쓸 수 있는 것은?",""],
            ["더 만들어보고 싶은 자동화는?",""]])],
  qs=["로봇 자동화가 늘어나면 우리 생활은 어떻게 바뀔까?",
      "이번 수업에서 새로 알게 된 가장 중요한 것은 무엇인가?"]))

# ============ 차시 렌더링 ============
for s in sessions:
    h(f"{s['no']}차시 학습활동지 — {s['title']}", size=16)
    sub("● 학습목표"); body(s['goal'])
    sub(f"● 사용 코드: {s['code']}")
    mission_box(s['mission'])
    for sec in s['tables']:
        sub(sec[0])
        table(sec[1], sec[2])
    sub("● 오늘의 정리 (질문에 답하기)")
    table(["질문","나의 답변"], [[q,""] for q in s['qs']])
    doc.add_page_break()

# ============ 최종 보고서 ============
h("아두이노 로봇팔 자동화 미션 보고서", size=16)
table(["항목","내용","항목","내용"],
      [["팀명","","작성일",""],
       ["사용 부품","로봇팔, 아두이노, 조이스틱, 서보모터","대표 미션",""],
       ["미션 동작 순서","","핵심 각도값",""]])
sub("● 자기평가표")
table(["평가 항목","예/아니오"],
      [["로봇팔의 관절 구조를 이해하였다","□ 예 / □ 아니오"],
       ["로봇팔을 안정적으로 조작할 수 있다","□ 예 / □ 아니오"],
       ["성공 자세를 각도값으로 기록할 수 있다","□ 예 / □ 아니오"],
       ["반복 동작을 함수로 만들 수 있다","□ 예 / □ 아니오"],
       ["명령/버튼으로 자동 동작을 실행할 수 있다","□ 예 / □ 아니오"],
       ["활동을 진로와 연결해 설명할 수 있다","□ 예 / □ 아니오"]])
doc.add_page_break()

# ============ 교사용 ============
h("교사용 — 관찰 기록 & 평가 루브릭", size=15, color="C00000")
sub("● 관찰 기록란")
table(["학생/팀","조작 숙련","코드 이해","함수/자동화","협력·태도","비고"],
      [["","","","","",""] for _ in range(4)])
sub("● 평가 루브릭")
table(["평가 영역","상","중","하"],
      [["로봇팔 구조 이해","서보 역할·연결을 정확히 설명","대체로 설명","설명을 어려워함"],
       ["코드 이해·수정","구조 파악 후 값을 의도대로 수정","일부 수정 가능","수정에 도움 필요"],
       ["함수·자동화","함수를 만들어 자동화 완성","예시 함수로 동작 구현","개념 이해 수준"],
       ["미션 수행","높은 성공률·창의적 설계","미션 성공","부분 성공"],
       ["진로 연결","로봇·자동화 진로와 구체적 연결","일반적 연결","연결 어려워함"]])
sub("● 생활기록부 활용 문장 예시")
for txt in [
 "아두이노 기반 6자유도 로봇팔의 관절 구조와 서보모터 제어 원리를 이해하고, 조이스틱·버튼 입력을 활용하여 자동 Pick & Place 동작을 구현함. 실패 원인을 각도값·그리퍼 제어·이동 속도 측면에서 분석·수정하는 등 공학적 문제해결 과정을 성실히 수행함.",
 "로봇팔의 기본 구조와 제어 방식을 이해하고, 물체 집기·이동 미션에 참여함. 각도값 기록과 함수 활용으로 동작을 자동화하는 과정을 경험하였으며, 활동을 로봇공학·스마트팩토리 진로와 연결하여 설명함.",
 "서보모터와 조이스틱 제어의 기초를 경험하고, 배선·코드 수정·각도 기록의 필요성을 이해함. 반복 실습을 통해 로봇 제어에 대한 관심을 확장함."]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
    r=p.add_run("• "+txt); r.font.size=Pt(9.5)

out="/home/user/myClaud/robot-arm-curriculum/worksheets/로봇팔_12차시_학습활동지_개선판.docx"
doc.save(out)
print("saved:", out)
